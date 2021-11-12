import flask
from flask import request, jsonify
from functrans import listToString,read_pdf,docx2txtrans,read_resumes,translator,skill_resumes,skills_doc,get_skjob_resum
#from préprocess import read_r,get_cleaned_words,remove_tags,html_parser,get_cleaned_wordsJ
import os
import io
import nltk
nltk.download('punkt')
nltk.download('stopwords')
import docx2txt
from docx import Document
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
import pandas as pd
import textwrap
from deep_translator import GoogleTranslator
import textwrap

import warnings
import os
warnings.filterwarnings("ignore")
# Import libraries
import sqlite3 as sql # sql connector
from datetime import datetime
import numpy as np             #for numerical computations like log,exp,sqrt etc
import pandas as pd            #for reading & storing data, pre-processing
import matplotlib.pylab as plt #for visualization
#import matplotlib.pyplot as plt
import re
from Cleaner import *
#from Cleaner import *
import tf_idf
exec(open("functrans.py").read())
exec(open("Cleaner.py").read())
from pyresparser import ResumeParser
import html.parser

import matplotlib.colors as mcolors
import gensim
import gensim.corpora as corpora
from gensim.summarization import keywords
from operator import index
from wordcloud import WordCloud
from pandas._config.config import options
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import Similar
from PIL import Image
import time


conn = sql.connect('data.sqlite')


# ## Construction CV utilisteur

user= pd.read_sql("""SELECT * FROM user ;""", conn)

user=user[['id_user','nom_user','prenom_user','ville','pays','telephone_user','email_user']]

formation_= pd.read_sql("""SELECT * FROM formation_ ;""", conn)

niveau_etude= pd.read_sql("""SELECT * FROM niveau_etude ;""", conn)


info_recherche= pd.read_sql("""SELECT * FROM info_recherche ;""", conn)

langue= pd.read_sql("""SELECT * FROM langue ;""", conn)
langue=langue.rename(columns = {'id': 'id_lang'}, inplace = False)

langue_speak= pd.read_sql("""SELECT * FROM langue_speak ;""", conn)


langues=pd.merge(langue, langue_speak[['id_lang','titre']],how= 'right' ,on="id_lang")

langues=langues[['id_lang','id_user','niveau','titre']]


experience_cv= pd.read_sql("""SELECT * FROM experience_cv ;""", conn)



competence_= pd.read_sql("""SELECT * FROM competence_ ;""", conn)

centreinteret= pd.read_sql("""SELECT * FROM centreinteret;""", conn)
centreinteret=centreinteret.rename(columns = {'candidat': 'id_user'}, inplace = False)

#recupération niveau d'étude en lettre
Niveaucv=pd.merge(info_recherche, niveau_etude[['id','nom_en']].rename(columns = {'id': 'niveau'}, inplace = False),how= 'left' ,on="niveau")


#Reconstruction CV
uuser=user.copy()
uuser.columns=['id_user', 'nom', 'prenom', 'ville', 'pays', 'telephone',
       'email']

Niveau_détude= Niveaucv[Niveaucv.columns.difference(['id_info', 'niveau', 'site_act', 'poste_rech', 'salaire',
       'type_salaire', 'ancienete', 'date_nat', 'lieu_nai', 'pays', 'ville',
       'deplacement', 'permis', 'status', 'dates', 'pays_search', 'modele_cv',
       'disponible', 'cv_physique', 'statut_cand', 'add_postale',
       'recherche_act', 'autre_ville', 'autre_diplome', 'autre_contrat',
       'precision_salaire', 'preselect_statut', 'device'])].rename(columns = {'nom_en': 'niveau'}, inplace = False)

Formation= formation_[formation_.columns.difference(['id','encours_fin','domaine','precise','statuts','type_forma','forma'])][['id_user','titre_forma','etab','description','lieu','mois1','annee1','mois2','annee2']]

Experience=experience_cv[experience_cv.columns.difference(['id','secteur_ent','status','encours_fin','contrat','autre_contrat'])][['id_user','poste','entreprise','description','lieu','mois1','annee1','mois2','annee2']]

Competence=competence_[competence_.columns.difference(['id','detail_comp','status','type_comp'])].rename(columns = {'description': 'compétence'}, inplace = False)

Langues=langues[langues.columns.difference(['id_lang'])].rename(columns = {'titre': 'langues'}, inplace = False)

Centreinteret=centreinteret[centreinteret.columns.difference(['id'])]



"""
cv1 = pd.merge(user, formation_ , how='left',on="id_user")
cv2=pd.merge(cv1, langues[['id_user','niveau','titre']] , how='left',on="id_user")
cv3=pd.merge(cv2, experience_cv , how='left',on="id_user")
cv4=pd.merge(cv3, competence_ , how='left',on="id_user")
df_cv=pd.merge(cv4, centreinteret, how='left',on="id_user")


# In[ ]:

dfcv1=df_cv[df_cv.columns]
df_cv['CV'] = dfcv1.astype(str).agg('-'.join, axis=1)
df_cv= df_cv.drop_duplicates(subset=['id_user'], keep='first')
df_cv = df_cv[(df_cv.CV!= ' ') & (df_cv.CV.notnull())]
"""

# ## Construction Job_despcription

# In[ ]:
entreprise_=pd.read_sql("""SELECT * FROM entreprise_ ;""", conn)

# In[ ]:
offre_= pd.read_sql("""SELECT * FROM offre_ ;""", conn)
# In[ ]:
offre_['Jobs_description']=offre_['nom']+offre_['description']+offre_['profil']
offre_=pd.merge(offre_, entreprise_[['id_ent','pays_ent']] , how='left',on="id_ent")
offre_= offre_.drop_duplicates(subset=['id'], keep='first')

# In[ ]:
candidature_= pd.read_sql("""SELECT * FROM candidature_ ;""", conn)


#####################------Modification

niveau_etude= pd.read_sql("""SELECT * FROM niveau_etude ;""", conn)
# attribution des score en tenant compte du nombre d'année d'étude
niveau_etude['score']=pd.DataFrame({'score':[6,10,13,14,15,16,17,18,19,20,0,10,15,15,12,12,12]})


experienceprofessionnelle= pd.read_sql("""SELECT * FROM experienceprofessionnelle ;""", conn)

experienceprofessionnelle['score']=pd.DataFrame({'score':[0,1,3,5,10,20,30]})

salaire= pd.read_sql("""SELECT * FROM salaire ;""", conn) 
salaire['score']=pd.DataFrame({'score':[5,10,20,30,40,50,60,70,80,90,5,10,20,30,40,50,5,10,20,30,40,50,0]})

type_remuneration= pd.read_sql("""SELECT * FROM type_remuneration ;""", conn)

pays= pd.read_sql("""SELECT * FROM pays ;""", conn) 


#Développons les fonctions pour calcuer ses critère secondaires
#######################  Niveau d' étude####################################################
def niveau(type_user,type_recrut,score_user,score_recruteur):
    if type_user==0 and type_recrut==0:
        if score_user >= score_recruteur:
            niveau_score="100%"
            #print("========================================100%")   
        if score_user < score_recruteur:
            niveau_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
            
    elif type_user==1 and type_recrut==1:
        if score_user >= score_recruteur:
            niveau_score="100%"
            #print("========================================100%")
        if score_user < score_recruteur:
            niveau_score=str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
    else:
        niveau_score="0%"
        
    return niveau_score
        
#######################  Expérience ####################################################
        
def exp(exp_user,exp_recruteur):
    if exp_user >= exp_recruteur:
        exp_score="100%"
        #print("========================================100%")
    elif exp_user < exp_recruteur:
        exp_score=str(exp_user/exp_recruteur*100)+"%"
        #print("=================="+str(exp_user/exp_recruteur*100)+"%")
        
    else:
        exp_score="0%"
    return exp_score
        
#######################  Salaire ####################################################
        
def sal(type_user,type_recrut,score_user,score_recruteur):
    if type_user==1 and type_recrut==1:
        if score_user == score_recruteur:
            sal_score= "100%"
            #print("=============================100%")
        if score_user > score_recruteur:
            sal_score= "100% >>>"
            #print("=================Supérieure à 100%")   
        if score_user < score_recruteur:
            sal_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
            
    elif type_user==2 and type_recrut==2:
        if score_user == score_recruteur:
            sal_score= "100%"
            #print("=============================100%")
        if score_user > score_recruteur:
            sal_score= "100% >>>"
            #print("=================Supérieure à 100%")   
        if score_user < score_recruteur:
            sal_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
            
    elif type_user==3 and type_recrut==3:
        if score_user == score_recruteur:
            sal_score= "100%"
            #print("=============================100%")
        if score_user > score_recruteur:
            sal_score= "100% >>>"
            #print("=================Supérieure à 100%")   
        if score_user < score_recruteur:
            sal_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
            
    elif type_user==4 and type_recrut==4:
        if score_user == score_recruteur:
            sal_score= "100%"
            #print("=============================100%")
        if score_user > score_recruteur:
            sal_score= "100% >>>"
            #print("=================Supérieure à 100%")   
        if score_user < score_recruteur:
            sal_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")

    elif type_user==5 and type_recrut==5:
        if score_user == score_recruteur:
            sal_score= "100%"
            #print("=============================100%")
        if score_user > score_recruteur:
            sal_score= "100% >>>"
            #print("=================Supérieure à 100%")   
        if score_user < score_recruteur:
            sal_score= str(score_user/score_recruteur*100)+"%"
            #print("=================="+str(score_user/score_recruteur*100)+"%")
    else:
        sal_score="0%"
    return sal_score


def process2 (id_candidat,id_ofcand):
    
    try:
        
        
        #Récupération des candidat qui ont postulé à l'offre
        #id_ofcand=str(45)

        cand_of=candidature_[candidature_["offre"].isin([id_ofcand])].reset_index()
        cand_of= cand_of.drop_duplicates(subset=['candidat'], keep='first')

        # Récupération des CV candidat qui ont postulé à l'offre
        CV=[]

        for i in list(cand_of['candidat']):
            CVV=[]
            id_candidatt=i
            
            User=uuser[uuser["id_user"].isin([id_candidatt])].reset_index()
            niveau_détude=Niveau_détude[Niveau_détude["id_user"].isin([id_candidatt])].reset_index()
            formation=Formation[Formation["id_user"].isin([id_candidatt])].reset_index()
            experience=Experience[Experience["id_user"].isin([id_candidatt])].reset_index()
            competence=Competence[Competence["id_user"].isin([id_candidatt])].reset_index()
            llangues=Langues[Langues["id_user"].isin([id_candidatt])].reset_index()
            ccentreinteret=Centreinteret[Centreinteret["id_user"].isin([id_candidatt])].reset_index()


            for key, value in User[['nom', 'prenom', 'ville', 'pays', 'telephone','email']].to_dict(orient="index")[0].items(): 
                cvv='%s:%s\n' % (key, value)
                CVV.append(cvv)

            for key, value in niveau_détude[['niveau']].to_dict(orient="index")[0].items(): 
                cvv='%s:%s\n' % (key, value)
                CVV.append(cvv)

            for i in range(0,formation.shape[0]):
                for key, value in formation.to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,experience.shape[0]):
                for key, value in experience.to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,competence.shape[0]):
                for key, value in competence[['compétence']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,llangues.shape[0]):
                for key, value in llangues[['langues']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,ccentreinteret.shape[0]):
                for key, value in ccentreinteret[['nom']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)
            CV.append(CVV)

        
        dfcv=cand_of.rename(columns = {'candidat': 'id_user'}, inplace = False)
        dfcv['CV']=pd.DataFrame(CV).astype(str).agg('-'.join, axis=1)
        CV_cand=dfcv[['id_user','CV']]

        # description de l'offre
        descriptall=offre_[offre_["id"].isin([id_ofcand])].reset_index() #recupération de l'offre
        descript= descriptall.drop_duplicates(subset=['Jobs_description'], keep='first') # pour recupéré la description sans duplique


        #création dun dossier portant l'id de l'offre
        newpath = id_ofcand 
        if not os.path.exists(newpath):
            os.makedirs(newpath)


        
        #Récupération des CV qui ont postulé à l'offre dans un document docx


        # Néttoyage du VC
        TAG_RE = re.compile(r'<[^>]+>')

        def remove_tags(text):
            return TAG_RE.sub('', text)

        def html_parser(text):
            return  html.parser.HTMLParser().unescape(text)




            
        documents=[]

        for i in range(0,CV_cand.shape[0]):
            CV=remove_tags(CV_cand['CV'][i])
            CV =html.parser.HTMLParser().unescape(CV)
            lines = textwrap.wrap(CV, 4999, break_long_words=False) #avoir les cv des candidat de l'offre et traduire
            CV=translator(lines)
            document = Document()
            document.add_paragraph(CV)
            fullText = []
            for para in document.paragraphs:
                fullText.append(para.text)
                data = '\n'.join(fullText)
                filename=str(CV_cand['id_user'][i])+'.docx'
                fileCV=[filename,data]
                documents.append(fileCV)
                

        # pour filtrer utilisateur dont le CV ne fera pas aumoins 36 mots
        docselect=[]
        for i in documents:
            S=i[1]
            K=S.split()
            if len(K)>=36:
                docselect.append(i)
        


        # néttoyage et transformation 



        def get_cleaned_words(document):
            for i in range(len(document)):
                raw = Cleaner(document[i][1])
                document[i].append(" ".join(raw[0]))
                document[i].append(" ".join(raw[1]))
                document[i].append(" ".join(raw[2]))
                sentence = tf_idf.do_tfidf(document[i][3].split(" "))
                document[i].append(sentence)
            return document

        Doc = get_cleaned_words(docselect)

        Database = pd.DataFrame(Doc, columns=[
                                "Noms", "Texte_Brute", "Texte_Traité", "Selection", "Selection_Reduite", "TF_Base"])


        #Faison appel à nos fonctions skill_resumes. Nous allons extraire les skills de l'enssemble des CV Candidats et les stocké dans une dataframe

        # script pour collecter les skills sur le document

        skdict=[]
        for i in range(0,Database.shape[0]):
            d=Database['Texte_Traité'][i]
            dict_skill=skills_doc(d)
            skdict.append(dict_skill)

        #itération pour recupérer les information  dans une liste simple
        listk=[]
        for i in skdict:
            n=i[0]
            listk.append(n)

        #recupération dans une dataframe    
        dataskill=get_skjob_resum(listk)


        # Néttoyage et traduction de l'offre



        job_trans=[]
        for i in range(0,descript.shape[0]):
            job=remove_tags(descript['Jobs_description'][i])
            job =html.parser.HTMLParser().unescape(job)
            lines = textwrap.wrap(job, 4999, break_long_words=False) #avoir les cv des candidat de l'offre et traduire
            job=translator(lines)
            job_trans.append(job)
            job=listToString(job_trans)

        #netoyage et reduction de l'offre
        def get_cleaned_wordsJ(document):
            for i in range(len(document)):
                raw = Cleaner(document[i][0])
                document[i].append(" ".join(raw[0]))
                document[i].append(" ".join(raw[1]))
                document[i].append(" ".join(raw[2]))
                sentence = tf_idf.do_tfidf(document[i][2].split(" "))
                document[i].append(sentence)
            return document

        Jd = get_cleaned_wordsJ([[job]])

        jd_database = pd.DataFrame(Jd, columns=[
                                "Texte_Brute", "Texte_Traité", "Selection", "Selection_Reduite", "TF_Base"])

        #Récupération des skills du Jobs description

        job_skill=get_skjob_resum(skills_doc(jd_database['Texte_Traité'][0]))


        #Nous recupérons nos données néttoyés

        Resumes=Database.copy()
        Jobs=jd_database.copy()


        # In[643]:


        #Notre fonction pour calculer le SCore de similarité CV_JOB

        def calculate_scores(resumes, job_description):
            scores = []
            for x in range(resumes.shape[0]):
                score = Similar.match(
                    resumes['TF_Base'][x], job_description['TF_Base'][0])
                scores.append(score)
            return scores


        # In[644]:


        Resumes['Scores'] = calculate_scores(Resumes, Jobs)

        Ranked_resumes = Resumes.sort_values(
            by=['Scores'], ascending=False).reset_index(drop=True)

        Ranked_resumes['Rangs'] = pd.DataFrame(
            [i for i in range(1, len(Ranked_resumes['Scores'])+1)])

        # recupération des donnée lié au rang du candidat
        lisdoc=[id_candidat+('.docx'),id_candidat+('.pdf'),id_candidat+('.doc')]

        # recupération du rang du candidat et score
        for i in list(Ranked_resumes['Noms']):
            if i in lisdoc:
                l=i
                Ranked_candresum=Ranked_resumes[Ranked_resumes["Noms"].isin([l])].reset_index()
                if int(Ranked_candresum.Rangs[0])<10:
                    rangs="You have in top ten applicant based on you Cv"
                    Ranked_CV= "You have scored"+' '+str(Ranked_candresum.Scores[0])+'\n'+ rangs
                else:
                    rangs="you are in"+' '+ str(Ranked_candresum.Rangs[0])+ 'position among the'+' '+str(len(Ranked_resumes))+' '+'candidates who applied'
                    Ranked_CV= "You have scored"+' '+str(Ranked_candresum.Scores[0])+'\n'+ rangs    
            else:
                Ranked_candresum="le candidat portant l'id"+ ' '+ id_candidat+ ' '+ "n'a pas postuler à cette offre"


        

        #Récupérons le top skill des candidats de l'offres
        #Calculons l’importance relative et la signification des Skills dans un document


        # In[647]:


        ############################################ TF-IDF Code ###################################


        # In[648]:


        def get_list_of_words(document):
            Document = []

            for a in document:
                raw = a.split(" ")
                Document.append(raw)

            return Document


        # In[649]:


        document = get_list_of_words(dataskill['Skills'])

        id2word = corpora.Dictionary(document)
        corpus = [id2word.doc2bow(text) for text in document]


        lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus, id2word=id2word, num_topics=1, random_state=100,
                                                    update_every=3, chunksize=100, passes=50, alpha='auto', per_word_topics=True)


        # In[650]:


        #Les Skills dominants


        # In[651]:


        #################################### LDA CODE ##############################################


        # In[652]:


        def format_topics_sentences(ldamodel, corpus):
            sent_topics_df = []
            for i, row_list in enumerate(ldamodel[corpus]):
                row = row_list[0] if ldamodel.per_word_topics else row_list
                row = sorted(row, key=lambda x: (x[1]), reverse=True)
                for j, (topic_num, prop_topic) in enumerate(row):
                    if j == 0:
                        wp = ldamodel.show_topic(topic_num)
                        topic_keywords = " ".join([word for word, prop in wp])
                        sent_topics_df.append(
                            [i, int(topic_num), round(prop_topic, 4)*100, topic_keywords])
                    else:
                        break

            return sent_topics_df


        # In[653]:


        ####################### CONFIGURATION DU CADRE DE DONNÉES POUR SUNBURST-GRAPH ############################


        # In[654]:


        df_topic_sents_keywords = format_topics_sentences(
            ldamodel=lda_model, corpus=corpus)
        df_some = pd.DataFrame(df_topic_sents_keywords, columns=[
                               'Document No', 'Skills dominant', 'Skills % contribution', 'Mot_clés'])
        df_some['Noms'] = Resumes['Noms']

        df = df_some

        #Comparaison et compte des skills candidat par rapports au Skills fréquament rencontrés chez la majorité

        countlst=[]
        interlst=[]
        globalskills=df.Mot_clés[0].split()
        for i in range(0,dataskill.shape[0]):
            loacalskills=dataskill['Skills'][i].split()
            inter=set(globalskills).intersection(loacalskills)
            interlst.append(inter)
            count = len(inter)
            countlst.append(count)
            
        #Creation de la dataframe Nom et skill

        comptskl = pd.DataFrame({'skills_e': interlst,'count': countlst})
        comptskl['Noms']=df['Noms']
        comptskl.head()


        stdict=[]
        # convert to string
        for i in range(0,comptskl.shape[0]):
            input= str(comptskl['skills_e'][i]).replace('{','').replace('}','').replace(",'","'")
            stdict.append(input)
            
        comptskl['skills_e'] = pd.DataFrame({'kills_e':stdict})
        #top skill
        sc_sf=comptskl[comptskl["Noms"].isin(lisdoc)].reset_index()
        Top_skill= "You have"+' '+str(sc_sf['count'][0])+' '+str(sc_sf.skills_e[0].split())+' '+"out of"+' '+str(len(df.Mot_clés[0].split()))+" top skills among all other applicants."+' '+str(df.Mot_clés[0].split())

        #Comparaisons Skill Job VS Skill candidat


        countlstJ=[]
        interlstJ=[]

        globalskills=job_skill.Skills[0].split() #skills job
        for i in range(0,dataskill.shape[0]):
            loacalskills=dataskill['Skills'][i].split()
            inter=set(globalskills).intersection(loacalskills)
            interlstJ.append(inter)
            count = len(inter)
            countlstJ.append(count)

        #Creation de la dataframe Nom et skill

        comptsklJ = pd.DataFrame({'skills_e': interlstJ,'count': countlstJ})
        comptsklJ['Noms']=df['Noms']
        comptsklJ.head()


        stdictJ=[]
        # convert to string
        for i in range(0,comptsklJ.shape[0]):
            input= str(comptsklJ['skills_e'][i]).replace('{','').replace('}','').replace(",'","'")
            stdictJ.append(input)

        comptsklJ['skills_e'] = pd.DataFrame({'kills_e':stdictJ})
        comptsklJ['score']=comptsklJ['count'].astype('int')/len(globalskills)*100 
        comptsklJ.sort_values(by=['score'], ascending=False,inplace=True)
        sj_sc=comptsklJ[comptsklJ["Noms"].isin(lisdoc)].reset_index()

        #skills
        Skills=str(int(sj_sc['count'][0])/len(job_skill.Skills[0].split())*100)+"%"

        # definition des critères secondaire candidat et Job
        critJob=descript[['nom','description','experience','niveau','type_salaire','pays_ent','salaire','device']]
        critcandidat=info_recherche[info_recherche["id_user"].isin([id_candidat])][['id_user','ancienete','niveau','type_salaire','pays_search','salaire','device']]
        critcandidat=critcandidat.reset_index(drop=True)
        critJob=critJob.reset_index(drop=True)

        # Recupération des critère au niveau des tables niveau_etude,experienceprofessionnelle,salaire

        # critère candidat
        crit_niv_cand=niveau_etude[niveau_etude["id"].isin([critcandidat["niveau"][0]])].reset_index(drop=True)
        crit_exp_cand=experienceprofessionnelle[experienceprofessionnelle["id"].isin([critcandidat["ancienete"][0]])].reset_index(drop=True)
        crit_sal_cand=critcandidat.copy()

        # critère Jobs
        crit_niv_job=niveau_etude[niveau_etude["id"].isin([critJob["niveau"][0]])].reset_index(drop=True)
        crit_exp_job=experienceprofessionnelle[experienceprofessionnelle["id"].isin([critJob["experience"][0]])].reset_index(drop=True)
        crit_sal_job=critJob.copy()

        # liste des candidat à retourner au recruteur basé sur CV
        rankcvuser=[]    
        for i in range(0,Ranked_resumes.head(10).shape[0]):
            recruCV=Ranked_resumes['Noms'][i]
            recruCV=recruCV.split('.')[0]
            rankcvuser.append(recruCV)
            
        ################CALCUL basé sur les fonction##############################""


        try:
            Niveau=niveau(int(crit_niv_cand['type'][0]),int(crit_niv_job['type'][0]),int(crit_niv_cand['score'][0]),int(crit_niv_job['score'][0]))
        except :
            Niveau="None"

        try:
            EExperience=exp(int(crit_exp_cand['score'][0]),int(crit_exp_job['score'][0]))
        except :
            EExperience="None"

            
        try:
            Salaire=sal(int(crit_sal_cand['type_salaire'][0]),int(crit_sal_job['type_salaire'][0]),int(crit_sal_cand['salaire'][0]),int(crit_sal_job['salaire'][0]),str(crit_sal_cand['device'][0]),str(crit_sal_job['device'][0]))
        except :
            Salaire="None"

        Results=[{'Ranked_CV':Ranked_CV,
                'Top_skill':Top_skill,
                'Skills':Skills,
                'Level':Niveau,
                'Experience':EExperience,
                'Salary':Salaire}
                ]
        return Results
    except :
        error="The result could not be found please complete your CV "
        return error




####################################"*******************RECRUTEUR"*****************#####################3
def process22 (id_ofcand):
    
    try:
        
        
        #Récupération des candidat qui ont postulé à l'offre
        #id_ofcand=str(45)

        cand_of=candidature_[candidature_["offre"].isin([id_ofcand])].reset_index()
        cand_of= cand_of.drop_duplicates(subset=['candidat'], keep='first')

        # Récupération des CV candidat qui ont postulé à l'offre
        CV=[]

        for i in list(cand_of['candidat']):
            CVV=[]
            id_candidatt=i
            
            User=uuser[uuser["id_user"].isin([id_candidatt])].reset_index()
            niveau_détude=Niveau_détude[Niveau_détude["id_user"].isin([id_candidatt])].reset_index()
            formation=Formation[Formation["id_user"].isin([id_candidatt])].reset_index()
            experience=Experience[Experience["id_user"].isin([id_candidatt])].reset_index()
            competence=Competence[Competence["id_user"].isin([id_candidatt])].reset_index()
            llangues=Langues[Langues["id_user"].isin([id_candidatt])].reset_index()
            ccentreinteret=Centreinteret[Centreinteret["id_user"].isin([id_candidatt])].reset_index()


            for key, value in User[['nom', 'prenom', 'ville', 'pays', 'telephone','email']].to_dict(orient="index")[0].items(): 
                cvv='%s:%s\n' % (key, value)
                CVV.append(cvv)

            for key, value in niveau_détude[['niveau']].to_dict(orient="index")[0].items(): 
                cvv='%s:%s\n' % (key, value)
                CVV.append(cvv)

            for i in range(0,formation.shape[0]):
                for key, value in formation.to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,experience.shape[0]):
                for key, value in experience.to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,competence.shape[0]):
                for key, value in competence[['compétence']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,llangues.shape[0]):
                for key, value in llangues[['langues']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)

            for i in range(0,ccentreinteret.shape[0]):
                for key, value in ccentreinteret[['nom']].to_dict(orient="index")[i].items(): 
                    cvv='%s:%s\n' % (key, value)
                    CVV.append(cvv)
            CV.append(CVV)

        
        dfcv=cand_of.rename(columns = {'candidat': 'id_user'}, inplace = False)
        dfcv['CV']=pd.DataFrame(CV).astype(str).agg('-'.join, axis=1)
        CV_cand=dfcv[['id_user','CV']]

        # description de l'offre
        descriptall=offre_[offre_["id"].isin([id_ofcand])].reset_index() #recupération de l'offre
        descript= descriptall.drop_duplicates(subset=['Jobs_description'], keep='first') # pour recupéré la description sans duplique


        #création dun dossier portant l'id de l'offre
        newpath = id_ofcand 
        if not os.path.exists(newpath):
            os.makedirs(newpath)


        
        #Récupération des CV qui ont postulé à l'offre dans un document docx


        # Néttoyage du VC
        TAG_RE = re.compile(r'<[^>]+>')

        def remove_tags(text):
            return TAG_RE.sub('', text)

        def html_parser(text):
            return  html.parser.HTMLParser().unescape(text)




            
        documents=[]

        for i in range(0,CV_cand.shape[0]):
            CV=remove_tags(CV_cand['CV'][i])
            CV =html.parser.HTMLParser().unescape(CV)
            lines = textwrap.wrap(CV, 4999, break_long_words=False) #avoir les cv des candidat de l'offre et traduire
            CV=translator(lines)
            document = Document()
            document.add_paragraph(CV)
            fullText = []
            for para in document.paragraphs:
                fullText.append(para.text)
                data = '\n'.join(fullText)
                filename=str(CV_cand['id_user'][i])+'.docx'
                fileCV=[filename,data]
                documents.append(fileCV)


        


        # pour filtrer utilisateur dont le CV ne fera pas aumoins 36 mots
        docselect=[]
        for i in documents:
            S=i[1]
            K=S.split()
            if len(K)>=36:
                docselect.append(i)
        


        # néttoyage et transformation 



        def get_cleaned_words(document):
            for i in range(len(document)):
                raw = Cleaner(document[i][1])
                document[i].append(" ".join(raw[0]))
                document[i].append(" ".join(raw[1]))
                document[i].append(" ".join(raw[2]))
                sentence = tf_idf.do_tfidf(document[i][3].split(" "))
                document[i].append(sentence)
            return document

        Doc = get_cleaned_words(docselect)

        Database = pd.DataFrame(Doc, columns=[
                                "Noms", "Texte_Brute", "Texte_Traité", "Selection", "Selection_Reduite", "TF_Base"])


        #Faison appel à nos fonctions skill_resumes. Nous allons extraire les skills de l'enssemble des CV Candidats et les stocké dans une dataframe

        # script pour collecter les skills sur le document

        skdict=[]
        for i in range(0,Database.shape[0]):
            d=Database['Texte_Traité'][i]
            dict_skill=skills_doc(d)
            skdict.append(dict_skill)

        #itération pour recupérer les information  dans une liste simple
        listk=[]
        for i in skdict:
            n=i[0]
            listk.append(n)

        #recupération dans une dataframe    
        dataskill=get_skjob_resum(listk)


        # Néttoyage et traduction de l'offre



        job_trans=[]
        for i in range(0,descript.shape[0]):
            job=remove_tags(descript['Jobs_description'][i])
            job =html.parser.HTMLParser().unescape(job)
            lines = textwrap.wrap(job, 4999, break_long_words=False) #avoir les cv des candidat de l'offre et traduire
            job=translator(lines)
            job_trans.append(job)
            job=listToString(job_trans)

        #netoyage et reduction de l'offre
        def get_cleaned_wordsJ(document):
            for i in range(len(document)):
                raw = Cleaner(document[i][0])
                document[i].append(" ".join(raw[0]))
                document[i].append(" ".join(raw[1]))
                document[i].append(" ".join(raw[2]))
                sentence = tf_idf.do_tfidf(document[i][2].split(" "))
                document[i].append(sentence)
            return document

        Jd = get_cleaned_wordsJ([[job]])

        jd_database = pd.DataFrame(Jd, columns=[
                                "Texte_Brute", "Texte_Traité", "Selection", "Selection_Reduite", "TF_Base"])

        #Récupération des skills du Jobs description

        job_skill=get_skjob_resum(skills_doc(jd_database['Texte_Traité'][0]))


        #Nous recupérons nos données néttoyés

        Resumes=Database.copy()
        Jobs=jd_database.copy()


        # In[643]:


        #Notre fonction pour calculer le SCore de similarité CV_JOB

        def calculate_scores(resumes, job_description):
            scores = []
            for x in range(resumes.shape[0]):
                score = Similar.match(
                    resumes['TF_Base'][x], job_description['TF_Base'][0])
                scores.append(score)
            return scores


        # In[644]:


        Resumes['Scores'] = calculate_scores(Resumes, Jobs)

        Ranked_resumes = Resumes.sort_values(
            by=['Scores'], ascending=False).reset_index(drop=True)

        Ranked_resumes['Rangs'] = pd.DataFrame(
            [i for i in range(1, len(Ranked_resumes['Scores'])+1)])


        

        #Récupérons le top skill des candidats de l'offres
        #Calculons l’importance relative et la signification des Skills dans un document


        # In[647]:


        ############################################ TF-IDF Code ###################################


        # In[648]:


        def get_list_of_words(document):
            Document = []

            for a in document:
                raw = a.split(" ")
                Document.append(raw)

            return Document


        # In[649]:


        document = get_list_of_words(dataskill['Skills'])

        id2word = corpora.Dictionary(document)
        corpus = [id2word.doc2bow(text) for text in document]


        lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus, id2word=id2word, num_topics=1, random_state=100,
                                                    update_every=3, chunksize=100, passes=50, alpha='auto', per_word_topics=True)


        # In[650]:


        #Les Skills dominants


        # In[651]:


        #################################### LDA CODE ##############################################


        # In[652]:


        def format_topics_sentences(ldamodel, corpus):
            sent_topics_df = []
            for i, row_list in enumerate(ldamodel[corpus]):
                row = row_list[0] if ldamodel.per_word_topics else row_list
                row = sorted(row, key=lambda x: (x[1]), reverse=True)
                for j, (topic_num, prop_topic) in enumerate(row):
                    if j == 0:
                        wp = ldamodel.show_topic(topic_num)
                        topic_keywords = " ".join([word for word, prop in wp])
                        sent_topics_df.append(
                            [i, int(topic_num), round(prop_topic, 4)*100, topic_keywords])
                    else:
                        break

            return sent_topics_df


        # In[653]:


        ####################### CONFIGURATION DU CADRE DE DONNÉES POUR SUNBURST-GRAPH ############################


        # In[654]:


        df_topic_sents_keywords = format_topics_sentences(
            ldamodel=lda_model, corpus=corpus)
        df_some = pd.DataFrame(df_topic_sents_keywords, columns=[
                               'Document No', 'Skills dominant', 'Skills % contribution', 'Mot_clés'])
        df_some['Noms'] = Resumes['Noms']

        df = df_some


        #Comparaisons Skill Job VS Skill candidat


        # In[660]:


        countlstJ=[]
        interlstJ=[]

        globalskills=job_skill.Skills[0].split() #skills job
        for i in range(0,dataskill.shape[0]):
            loacalskills=dataskill['Skills'][i].split()
            inter=set(globalskills).intersection(loacalskills)
            interlstJ.append(inter)
            count = len(inter)
            countlstJ.append(count)

        #Creation de la dataframe Nom et skill

        comptsklJ = pd.DataFrame({'skills_e': interlstJ,'count': countlstJ})
        comptsklJ['Noms']=df['Noms']
        comptsklJ.head()


        stdictJ=[]
        # convert to string
        for i in range(0,comptsklJ.shape[0]):
            input= str(comptsklJ['skills_e'][i]).replace('{','').replace('}','').replace(",'","'")
            stdictJ.append(input)

        comptsklJ['skills_e'] = pd.DataFrame({'kills_e':stdictJ})
        comptsklJ['score']=comptsklJ['count'].astype('int')/len(globalskills)*100 
        comptsklJ.sort_values(by=['score'], ascending=False,inplace=True)

        # liste des candidat à retourner au recruteur basé sur CV
        rankcvuser=[]    
        for i in range(0,Ranked_resumes.head(10).shape[0]):
            recruCV=Ranked_resumes['Noms'][i]
            recruCV=recruCV.split('.')[0]
            rankcvuser.append(recruCV)
            
        #liste des candidat à retourner au recruteur basé sur CV
        rankskilluser=[]    
        for i in range(0,comptsklJ.head(10).shape[0]):
            recruSK=comptsklJ['Noms'][i]
            recruSK=recruSK.split('.')[0]
            rankskilluser.append(recruSK)
            
        Resultrecruteur=[{'Ranked_CV_id': 'Top ten user based on CV'+':'+' '+str(rankcvuser),
                'Skillsuser_id':'Top ten user based on Skill'+':'+' '+str(rankskilluser)
                }
                ]
        return Resultrecruteur
    except :
        error="The result could not be found"
        return error


############################################*************API**********####################################""


app = flask.Flask(__name__)

@app.route('/', methods=['GET'])
def home():
    return '''<h1>Distant Jobs Matching from Sawajob</h1>
<p>A prototype API for distant Match Jobs CV.</p>'''

@app.errorhandler(404)
def page_not_found(e):
    return "<h1>404</h1><p>The resource could not be found.</p>", 404
#http://127.0.0.1:5000/predict_api?id1=133&id2=45


@app.route('/candidat',methods=['POST','GET'])
def predict_cand():
    id_candidat = str(request.args.get('id1'))
    id_ofcand = str(request.args.get('id2'))
    return jsonify(process2 (id_candidat,id_ofcand))

@app.route('/recruteur',methods=['POST','GET'])
def predict_recru():
    id_ofcand = str(request.args.get('id2'))
    return jsonify(process22 (id_ofcand))

# We only need this for local development.
if __name__ == '__main__':
    app.run(host='0.0.0.0',port=8080)


